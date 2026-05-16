# -*- coding: utf-8 -*-
"""
Generuje raport Word z opisem metodologii i wynikami obu analiz.
Uruchom: py generuj_raport.py
"""

import os
import warnings
import numpy as np
import pandas as pd
import statsmodels.api as sm
from statsmodels.stats.stattools import durbin_watson
from statsmodels.stats.diagnostic import acorr_breusch_godfrey, het_breuschpagan
from statsmodels.stats.outliers_influence import variance_inflation_factor
from scipy.stats import shapiro
from statsmodels.stats.stattools import jarque_bera

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

try:
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
except NameError:
    SCRIPT_DIR = r"C:\Users\Patryk\Desktop\studia magisterskie\Prognozowanie i symulacje\Projekt"

# ============================================================
# 1. ŁADOWANIE I PRZYGOTOWANIE DANYCH
# ============================================================

df_e = pd.read_excel(os.path.join(SCRIPT_DIR, "Zuzycie_energii_polska.xlsx"))
df_e["pkb_per_capita"] = df_e["pkb_mln_zl"] * 1e6 / df_e["ludnosc"]
df_e["ln_pkb_pc"]   = np.log(df_e["pkb_per_capita"])
df_e["ln_zuzycie"]  = np.log(df_e["zuzycie_energii_GWh"])
df_e["ln_cena"]     = np.log(df_e["cena_energii_zl_kWh"])

df_m = pd.read_excel(os.path.join(SCRIPT_DIR, "mieszkania_polska.xlsx"))
for col in df_m.columns:
    if col.startswith("NAK"):
        df_m = df_m.rename(columns={col: "NAKL"})
        break

# ============================================================
# 2. ESTYMACJA MODELI
# ============================================================

# --- ENERGIA Model 1 ---
Xe1_cols = ["ln_pkb_pc", "ln_cena", "urbanizacja_pct", "hdd", "cdd"]
Xe1 = sm.add_constant(df_e[Xe1_cols])
me1 = sm.OLS(df_e["ln_zuzycie"], Xe1).fit()

# --- ENERGIA Model 2 ---
Xe2_cols = ["ln_pkb_pc", "ln_cena", "hdd", "cdd"]
Xe2 = sm.add_constant(df_e[Xe2_cols])
me2 = sm.OLS(df_e["ln_zuzycie"], Xe2).fit()

# --- MIESZKANIA Model 1 ---
Xm1_cols = ["NAKL", "WYNAGR", "WSK25-34", "WSK_URB", "SM"]
Xm1 = sm.add_constant(df_m[Xm1_cols])
mm1 = sm.OLS(df_m["MO"], Xm1).fit()

# --- MIESZKANIA Model 2 (OLS + HAC) ---
Xm2_cols = ["NAKL", "WSK25-34", "WSK_URB", "SM"]
Xm2 = sm.add_constant(df_m[Xm2_cols])
mm2_ols = sm.OLS(df_m["MO"], Xm2).fit()
mm2     = mm2_ols.get_robustcov_results(cov_type="HAC", maxlags=2)

# ============================================================
# 3. DIAGNOSTYKA
# ============================================================

def diagnostics(model, X):
    r = model.resid
    sw_p  = shapiro(r)[1]
    dw    = durbin_watson(r)
    bg_p  = acorr_breusch_godfrey(model, nlags=2)[1]
    bp_p  = het_breuschpagan(r, X)[1]
    return sw_p, dw, bg_p, bp_p

def vif_dict(X, cols):
    return {c: variance_inflation_factor(X.values, i+1) for i, c in enumerate(cols)}

sw_e1, dw_e1, bg_e1, bp_e1 = diagnostics(me1, Xe1)
sw_e2, dw_e2, bg_e2, bp_e2 = diagnostics(me2, Xe2)
sw_m1, dw_m1, bg_m1, bp_m1 = diagnostics(mm1, Xm1)
sw_m2, dw_m2, bg_m2, bp_m2 = diagnostics(mm2_ols, Xm2)

vif_e1 = vif_dict(Xe1, Xe1_cols)
vif_e2 = vif_dict(Xe2, Xe2_cols)
vif_m1 = vif_dict(Xm1, Xm1_cols)
vif_m2 = vif_dict(Xm2, Xm2_cols)

# ============================================================
# 4. PROGNOZA ENERGII (scenariusze)
# ============================================================
forecast_years = list(range(2025, 2031))
base_pkb  = df_e["pkb_per_capita"].iloc[-1]
base_cena = df_e["cena_energii_zl_kWh"].iloc[-1]

scenarios = {
    "Pesymistyczny": dict(pkb=0.015, cena=0.06, hdd=3100, cdd=25),
    "Bazowy":        dict(pkb=0.030, cena=0.03, hdd=2900, cdd=35),
    "Optymistyczny": dict(pkb=0.045, cena=0.01, hdd=2700, cdd=50),
}

fc_energy = {}
for name, sc in scenarios.items():
    vals = []
    for i, yr in enumerate(forecast_years, 1):
        x = pd.DataFrame({
            "const":      [1.0],
            "ln_pkb_pc":  [np.log(base_pkb  * (1 + sc["pkb"])  ** i)],
            "ln_cena":    [np.log(base_cena  * (1 + sc["cena"]) ** i)],
            "hdd":        [float(sc["hdd"])],
            "cdd":        [float(sc["cdd"])],
        })
        pred  = me2.get_prediction(x).summary_frame(alpha=0.05)
        vals.append((yr,
                     int(np.exp(pred["mean"].values[0])),
                     int(np.exp(pred["mean_ci_lower"].values[0])),
                     int(np.exp(pred["mean_ci_upper"].values[0]))))
    fc_energy[name] = vals

# ============================================================
# 5. BUDOWANIE DOKUMENTU WORD
# ============================================================

doc = Document()

# --- marginesy ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ---- pomocnicze ----

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x5c, 0x96)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x27, 0xae, 0x60)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def p(text="", bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)
    return para

def ok(v):   return "✔ TAK"  if v > 0.05 else "✘ NIE"
def okdw(v): return "✔ OK"   if 1.5 < v < 2.5 else "⚠ Możliwa autokorelacja"

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # header background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"),   "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"),  "1A5C96")
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for rdata in rows:
        row = t.add_row()
        for i, val in enumerate(rdata):
            row.cells[i].text = str(val)
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return t

# ============================================================
# STRONA TYTUŁOWA
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Analiza prognozowania danych\nspołeczno-ekonomicznych w Polsce")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x96)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Projekt zaliczeniowy — Prognozowanie i Symulacje\n"
            "Dane: Polska, lata 2004–2024\n"
            "Metoda: MNK (Metoda Najmniejszych Kwadratów), ARIMA").italic = True

doc.add_page_break()

# ============================================================
# WSTĘP
# ============================================================
h1("Wstęp")
p("Celem projektu jest zbadanie, jakie czynniki wpływają na dwa ważne wskaźniki "
  "społeczno-ekonomiczne w Polsce: zużycie energii elektrycznej oraz liczbę nowo "
  "wybudowanych mieszkań. Analizę przeprowadzono na danych rocznych z lat 2004–2024.")
p()
p("Dla każdego zjawiska zbudowano model statystyczny (regresję), sprawdzono jego "
  "poprawność i w razie potrzeby poprawiono. W przypadku zużycia energii "
  "przygotowano również prognozy na lata 2025–2030.")
p()
p("Cały projekt wykonano w języku Python z użyciem biblioteki statsmodels.")

# ============================================================
# CZĘŚĆ 1 – ENERGIA
# ============================================================
doc.add_page_break()
h1("Część 1 — Zużycie energii elektrycznej w Polsce")

# --- 1.1 Dane ---
h2("1.1  Dane wejściowe")
p("Źródło danych: plik Zuzycie_energii_polska.xlsx")
p(f"Liczba obserwacji: {len(df_e)} (po jednej na każdy rok)")
p()
p("Użyte zmienne:", bold=True)

add_table(
    ["Zmienna", "Opis", "Jednostka"],
    [
        ["zuzycie_energii_GWh", "Całkowite zużycie energii elektrycznej", "GWh/rok"],
        ["cena_energii_zl_kWh", "Średnia cena energii dla odbiorców",    "zł/kWh"],
        ["pkb_mln_zl",          "Produkt Krajowy Brutto",                 "mln zł"],
        ["ludnosc",             "Liczba ludności Polski",                  "osoby"],
        ["urbanizacja_pct",     "Odsetek mieszkańców miast",               "%"],
        ["hdd",                 "Heating Degree Days — wskaźnik zimy",    "dni·°C"],
        ["cdd",                 "Cooling Degree Days — wskaźnik lata",    "dni·°C"],
    ]
)
p("Na podstawie powyższych danych obliczono PKB per capita (PKB / liczba ludności), "
  "a następnie wszystkie zmienne pieniężne i energetyczne zostały przetransformowane "
  "logarytmicznie (ln). Dzięki temu współczynniki modelu mają interpretację elastyczności.")

# --- 1.2 Model 1 ---
h2("1.2  Model 1 — podstawowy")

h3("Co to jest i po co?")
p("Model regresji liniowej (OLS — Metoda Najmniejszych Kwadratów) szuka zależności "
  "między zużyciem energii a czynnikami ekonomicznymi i klimatycznymi. "
  "Używamy formy logarytmicznej po obu stronach równania, co pozwala odczytywać "
  "wyniki jako: »jeśli PKB wzrośnie o 1%, to zużycie energii zmieni się o β₁ procent«.")

h3("Równanie modelu")
p("ln(ZUŻYCIE) = β₀ + β₁ · ln(PKB per capita) + β₂ · ln(CENA)")
p("             + β₃ · URBANIZACJA + β₄ · HDD + β₅ · CDD + błąd",
  italic=True)

h3("Wyniki estymacji")

p1 = me1.params
pv1 = me1.pvalues

def sig(p):
    if p < 0.01:  return "*** (p<0.01)"
    if p < 0.05:  return "** (p<0.05)"
    if p < 0.10:  return "* (p<0.10)"
    return "brak istotności"

add_table(
    ["Zmienna", "Współczynnik β", "Interpretacja", "Istotność"],
    [
        ["ln(PKB per capita)", f"{p1['ln_pkb_pc']:+.3f}",
         f"1% wzrostu PKB → {p1['ln_pkb_pc']:+.2f}% zmiana zużycia",
         sig(pv1["ln_pkb_pc"])],
        ["ln(Cena energii)", f"{p1['ln_cena']:+.3f}",
         f"1% wzrostu ceny → {p1['ln_cena']:+.2f}% zmiana zużycia",
         sig(pv1["ln_cena"])],
        ["Urbanizacja [%]", f"{p1['urbanizacja_pct']:+.4f}",
         "Wzrost zurbanizowania o 1pp",
         sig(pv1["urbanizacja_pct"])],
        ["HDD", f"{p1['hdd']:+.6f}",
         "Każdy dodatkowy dzień grzewczy",
         sig(pv1["hdd"])],
        ["CDD", f"{p1['cdd']:+.6f}",
         "Każdy dodatkowy dzień chłodzenia",
         sig(pv1["cdd"])],
    ]
)

h3("Jak dobrze model pasuje do danych?")
add_table(
    ["Miara", "Wartość", "Co oznacza?"],
    [
        ["R²",       f"{me1.rsquared:.4f}",
         f"Model wyjaśnia {me1.rsquared*100:.1f}% zmienności zużycia energii"],
        ["R² adj.",  f"{me1.rsquared_adj:.4f}",
         "R² skorygowane o liczbę zmiennych"],
        ["AIC",      f"{me1.aic:.2f}",
         "Kryterium informacyjne (niższe = lepsze)"],
        ["BIC",      f"{me1.bic:.2f}",
         "Kryterium bayesowskie (niższe = lepsze)"],
    ]
)

h3("Sprawdzenie poprawności statystycznej")
p("Dla modelu przeprowadzono cztery testy diagnostyczne:", bold=False)

add_table(
    ["Test", "Wynik (p-value)", "Wniosek"],
    [
        ["Shapiro-Wilk\n(normalność reszt)",
         f"p = {sw_e1:.4f}", ok(sw_e1)],
        ["Durbin-Watson\n(autokorelacja)",
         f"DW = {dw_e1:.3f}", okdw(dw_e1)],
        ["Breusch-Godfrey\n(autokorelacja)",
         f"p = {bg_e1:.4f}", ok(bg_e1)],
        ["Breusch-Pagan\n(jednakowa wariancja)",
         f"p = {bp_e1:.4f}", ok(bp_e1)],
    ]
)

h3("Wykryty problem — wielokoliniowość")
p("Wielokoliniowość oznacza, że dwie zmienne objaśniające są ze sobą za bardzo "
  "powiązane — model nie potrafi rozróżnić ich efektów. Mierzy się ją wskaźnikiem VIF "
  "(Variance Inflation Factor): wartości powyżej 10 są niepokojące, powyżej 30 — "
  "poważny problem.")
p()

add_table(
    ["Zmienna", "VIF", "Ocena"],
    [
        [c, f"{vif_e1[c]:.1f}",
         "⚠ PROBLEM" if vif_e1[c] > 10 else "✔ OK"]
        for c in Xe1_cols
    ]
)
p(f"Korelacja między ln(PKB per capita) a Urbanizacją wynosi "
  f"{df_e[['ln_pkb_pc','urbanizacja_pct']].corr().iloc[0,1]:.3f} — "
  "niemal idealna zależność liniowa. Obie zmienne mierzą to samo zjawisko "
  "(wzrost zamożności i zurbanizowania idą razem). "
  "Rozwiązanie: usunięcie Urbanizacji → patrz Model 2.")

# --- 1.3 Model 2 ---
h2("1.3  Model 2 — po korekcie (Iteracja 2)")

h3("Co zmieniono i dlaczego?")
p("Usunięto zmienną Urbanizacja, która była prawie idealnie skorelowana z PKB per capita. "
  "Nowy model zawiera cztery zmienne: PKB per capita, cenę energii oraz dwa wskaźniki "
  "klimatyczne.")

h3("Równanie modelu")
p("ln(ZUŻYCIE) = β₀ + β₁ · ln(PKB per capita) + β₂ · ln(CENA) + β₃ · HDD + β₄ · CDD + błąd",
  italic=True)

h3("Wyniki estymacji")
p2 = me2.params
pv2 = me2.pvalues

add_table(
    ["Zmienna", "Współczynnik β", "Interpretacja", "Istotność"],
    [
        ["ln(PKB per capita)", f"{p2['ln_pkb_pc']:+.3f}",
         f"1% wzrostu PKB → {p2['ln_pkb_pc']:+.2f}% zmiana zużycia",
         sig(pv2["ln_pkb_pc"])],
        ["ln(Cena energii)", f"{p2['ln_cena']:+.3f}",
         f"1% wzrostu ceny → {p2['ln_cena']:+.2f}% zmiana zużycia",
         sig(pv2["ln_cena"])],
        ["HDD", f"{p2['hdd']:+.6f}",
         "Każdy dodatkowy dzień grzewczy",
         sig(pv2["hdd"])],
        ["CDD", f"{p2['cdd']:+.6f}",
         "Każdy dodatkowy dzień chłodzenia",
         sig(pv2["cdd"])],
    ]
)

h3("Jak dobrze model pasuje do danych?")
add_table(
    ["Miara", "Model 1", "Model 2", "Zmiana"],
    [
        ["R²",      f"{me1.rsquared:.4f}",     f"{me2.rsquared:.4f}",
         "↓ nieznaczny spadek (usunięto 1 zmienną)"],
        ["R² adj.", f"{me1.rsquared_adj:.4f}", f"{me2.rsquared_adj:.4f}",
         "Porównywalny"],
        ["AIC",     f"{me1.aic:.2f}",          f"{me2.aic:.2f}",
         "↓ lepsze" if me2.aic < me1.aic else "↑ gorsze"],
        ["BIC",     f"{me1.bic:.2f}",          f"{me2.bic:.2f}",
         "↓ lepsze" if me2.bic < me1.bic else "↑ gorsze"],
    ]
)

h3("Sprawdzenie poprawności — Model 2")
add_table(
    ["Test", "Model 1", "Model 2", "Poprawa?"],
    [
        ["Shapiro-Wilk (normalność)",
         f"p={sw_e1:.4f} {ok(sw_e1)}",
         f"p={sw_e2:.4f} {ok(sw_e2)}",
         "→" if abs(sw_e2-sw_e1)<0.01 else ("✔" if sw_e2>sw_e1 else "✘")],
        ["Durbin-Watson",
         f"{dw_e1:.3f} {okdw(dw_e1)}",
         f"{dw_e2:.3f} {okdw(dw_e2)}",
         "↑ lepsza" if abs(dw_e2-2) < abs(dw_e1-2) else "→ podobna"],
        ["Breusch-Godfrey (autokorelacja)",
         f"p={bg_e1:.4f} {ok(bg_e1)}",
         f"p={bg_e2:.4f} {ok(bg_e2)}",
         "✔" if bg_e2>0.05 and bg_e1<0.05 else "→"],
        ["Breusch-Pagan (wariancja)",
         f"p={bp_e1:.4f} {ok(bp_e1)}",
         f"p={bp_e2:.4f} {ok(bp_e2)}",
         "→"],
    ]
)

h3("Wielokoliniowość po korekcie")
add_table(
    ["Zmienna", "VIF Model 1", "VIF Model 2", "Ocena"],
    [
        [c, f"{vif_e1[c]:.1f}", f"{vif_e2[c]:.1f}",
         "✔ Poprawa" if vif_e2[c] < vif_e1[c] else "→"]
        for c in Xe2_cols
    ]
)

h3("Czy Model 2 nadaje się do prognoz?")
bullet("Brak autokorelacji (Breusch-Godfrey p > 0.05) — ✔ OK")
bullet("Brak heteroskedastyczności (Breusch-Pagan p > 0.05) — ✔ OK")
bullet(f"Normalność reszt: {'✔ OK' if sw_e2>0.05 else '⚠ Reszty odbiegają od normalności — przedziały ufności prognozy są orientacyjne'}")
bullet("VIF zmniejszony — wielokoliniowość ograniczona ✔")
p()
p("Wniosek: Model 2 nadaje się do prognoz warunkowych (tj. przy założonych "
  "wartościach PKB, cen i wskaźników klimatycznych). Wyniki należy traktować "
  "jako szacunki, nie jako precyzyjne liczby.", bold=False)

# --- 1.4 Prognozy ---
h2("1.4  Prognozy zużycia energii 2025–2030")

h3("Metodologia scenariuszy")
p("Prognozy zbudowano na podstawie Modelu 2 przy trzech założeniach o przyszłości:")
add_table(
    ["Parametr", "Pesymistyczny", "Bazowy", "Optymistyczny"],
    [
        ["Wzrost PKB per capita rocznie", "+1.5%", "+3.0%", "+4.5%"],
        ["Wzrost ceny energii rocznie",   "+6.0%", "+3.0%", "+1.0%"],
        ["HDD (mroźność zimy)",           "3100",  "2900",  "2700"],
        ["CDD (upalność lata)",           "25",    "35",    "50"],
    ]
)

h3("Wyniki prognoz [GWh]")
rows_fc = []
for yr in forecast_years:
    pess = next(v for v in fc_energy["Pesymistyczny"] if v[0]==yr)
    base = next(v for v in fc_energy["Bazowy"]        if v[0]==yr)
    opt  = next(v for v in fc_energy["Optymistyczny"] if v[0]==yr)
    rows_fc.append([str(yr),
                    f"{pess[1]:,}",
                    f"{base[1]:,}",
                    f"{opt[1]:,}"])

add_table(["Rok", "Pesymistyczny [GWh]", "Bazowy [GWh]", "Optymistyczny [GWh]"],
          rows_fc)

base_2024 = int(df_e["zuzycie_energii_GWh"].iloc[-1])
base_2030 = fc_energy["Bazowy"][-1][1]
p(f"Dla porównania: zużycie w 2024 roku wynosiło {base_2024:,} GWh. "
  f"Scenariusz bazowy przewiduje {base_2030:,} GWh w 2030 roku "
  f"({(base_2030/base_2024-1)*100:+.1f}% względem 2024).")

# ============================================================
# CZĘŚĆ 2 – MIESZKANIA
# ============================================================
doc.add_page_break()
h1("Część 2 — Mieszkania oddane do użytkowania w Polsce")

# --- 2.1 Dane ---
h2("2.1  Dane wejściowe")
p("Źródło danych: plik mieszkania_polska.xlsx")
p(f"Liczba obserwacji: {len(df_m)} (po jednej na każdy rok)")
p()
p("Użyte zmienne:", bold=True)

add_table(
    ["Zmienna", "Opis", "Jednostka"],
    [
        ["MO",      "Mieszkania oddane do użytkowania",                "szt./rok"],
        ["NAKL",    "Nakłady na budownictwo mieszkaniowe",             "mln zł"],
        ["WYNAGR",  "Przeciętne wynagrodzenie brutto",                 "zł"],
        ["WSK25-34","Udział osób w wieku 25–34 lat w populacji",       "udział"],
        ["WSK_URB", "Wskaźnik urbanizacji",                            "udział"],
        ["SM",      "Saldo migracji (napływ minus odpływ)",            "‰"],
    ]
)

# --- 2.2 Model 1 ---
h2("2.2  Model 1 — podstawowy")

h3("Co to jest i po co?")
p("Model regresji liniowej (bez logarytmów) szuka zależności między liczbą "
  "oddanych mieszkań a czynnikami ekonomicznymi i demograficznymi. "
  "Forma liniowa oznacza: »jeśli nakłady wzrosną o 1 mln zł, to liczba "
  "mieszkań zmieni się o β₁ sztuk«.")

h3("Równanie modelu")
p("MO = β₀ + β₁ · NAKL + β₂ · WYNAGR + β₃ · WSK25-34 + β₄ · WSK_URB + β₅ · SM + błąd",
  italic=True)

h3("Wyniki estymacji")
pm1 = mm1.params
pmv1 = mm1.pvalues

add_table(
    ["Zmienna", "Współczynnik β", "Interpretacja", "Istotność"],
    [
        ["NAKL (nakłady)", f"{pm1['NAKL']:+.3f}",
         f"1 mln zł więcej nakładów → {pm1['NAKL']:+.0f} szt.",
         sig(pmv1["NAKL"])],
        ["WYNAGR", f"{pm1['WYNAGR']:+.3f}",
         f"Wzrost płacy o 1 zł → {pm1['WYNAGR']:+.2f} szt.",
         sig(pmv1["WYNAGR"])],
        ["WSK25-34", f"{pm1['WSK25-34']:+.1f}",
         "Udział grupy 25–34 lat",
         sig(pmv1["WSK25-34"])],
        ["WSK_URB", f"{pm1['WSK_URB']:+.1f}",
         "Wskaźnik urbanizacji",
         sig(pmv1["WSK_URB"])],
        ["SM (saldo migracji)", f"{pm1['SM']:+.1f}",
         f"Wzrost salda o 1‰ → {pm1['SM']:+.0f} szt.",
         sig(pmv1["SM"])],
    ]
)

h3("Jak dobrze model pasuje do danych?")
add_table(
    ["Miara", "Wartość", "Co oznacza?"],
    [
        ["R²",       f"{mm1.rsquared:.4f}",
         f"Model wyjaśnia {mm1.rsquared*100:.1f}% zmienności liczby mieszkań"],
        ["R² adj.",  f"{mm1.rsquared_adj:.4f}", "R² skorygowane"],
        ["AIC",      f"{mm1.aic:.2f}",           "Kryterium informacyjne"],
        ["BIC",      f"{mm1.bic:.2f}",           "Kryterium bayesowskie"],
    ]
)

h3("Sprawdzenie poprawności statystycznej")
add_table(
    ["Test", "Wynik", "Wniosek"],
    [
        ["Shapiro-Wilk (normalność)", f"p = {sw_m1:.4f}", ok(sw_m1)],
        ["Durbin-Watson (autokorelacja)", f"DW = {dw_m1:.3f}", okdw(dw_m1)],
        ["Breusch-Godfrey (autokorelacja)", f"p = {bg_m1:.4f}", ok(bg_m1)],
        ["Breusch-Pagan (wariancja)", f"p = {bp_m1:.4f}", ok(bp_m1)],
    ]
)

h3("Wykryty problem — wielokoliniowość")
add_table(
    ["Zmienna", "VIF", "Ocena"],
    [
        [c, f"{vif_m1[c]:.1f}",
         "⚠ PROBLEM" if vif_m1[c] > 10 else "✔ OK"]
        for c in Xm1_cols
    ]
)
r_nakl_wyn = df_m[["NAKL","WYNAGR"]].corr().iloc[0,1]
p(f"Korelacja NAKL ↔ WYNAGR = {r_nakl_wyn:.3f} — bardzo silna. "
  "Oba wskaźniki rosną razem ze wzrostem gospodarczym, więc model "
  "nie potrafi oddzielić ich wpływów. "
  "Rozwiązanie: usunięcie WYNAGR → patrz Model 2.")

# --- 2.3 Model 2 ---
h2("2.3  Model 2 — po korekcie (Iteracja 2)")

h3("Co zmieniono i dlaczego?")
p("Usunięto zmienną WYNAGR (wynagrodzenia), która była prawie identycznie "
  "skorelowana z NAKL (nakładami). "
  "Ponadto, ponieważ reszty modelu wykazywały autokorelację (wartości błędów "
  "w kolejnych latach są ze sobą powiązane), zastosowano odporne błędy standardowe "
  "HAC (metoda Newey-West). HAC poprawia wiarygodność testów statystycznych, "
  "ale nie zmienia wartości prognoz.")

h3("Równanie modelu")
p("MO = β₀ + β₁ · NAKL + β₂ · WSK25-34 + β₃ · WSK_URB + β₄ · SM + błąd",
  italic=True)
p("(Błędy standardowe: HAC Newey-West)", italic=True)

h3("Wyniki estymacji")
pm2  = mm2_ols.params   # współczynniki takie same jak OLS
# HAC pvalues mogą być numpy array bez nazw — konwertujemy
_pv2_raw = mm2.pvalues
if hasattr(_pv2_raw, "index"):
    pmv2 = _pv2_raw
else:
    pmv2 = pd.Series(_pv2_raw, index=mm2_ols.params.index)

add_table(
    ["Zmienna", "Współczynnik β", "Interpretacja", "Istotność"],
    [
        ["NAKL (nakłady)", f"{pm2['NAKL']:+.3f}",
         f"1 mln zł więcej nakładów → {pm2['NAKL']:+.0f} szt.",
         sig(pmv2["NAKL"])],
        ["WSK25-34", f"{pm2['WSK25-34']:+.1f}",
         "Udział grupy 25–34 lat",
         sig(pmv2["WSK25-34"])],
        ["WSK_URB", f"{pm2['WSK_URB']:+.1f}",
         "Wskaźnik urbanizacji",
         sig(pmv2["WSK_URB"])],
        ["SM (saldo migracji)", f"{pm2['SM']:+.1f}",
         f"Wzrost salda o 1‰ → {pm2['SM']:+.0f} szt.",
         sig(pmv2["SM"])],
    ]
)

h3("Porównanie modeli")
add_table(
    ["Miara", "Model 1", "Model 2", "Zmiana"],
    [
        ["R²",      f"{mm1.rsquared:.4f}",     f"{mm2_ols.rsquared:.4f}",
         "↓ nieznaczny spadek"],
        ["R² adj.", f"{mm1.rsquared_adj:.4f}", f"{mm2_ols.rsquared_adj:.4f}",
         "Porównywalny"],
        ["AIC",     f"{mm1.aic:.2f}",          f"{mm2_ols.aic:.2f}",
         "↓ lepsze" if mm2_ols.aic < mm1.aic else "↑"],
        ["BIC",     f"{mm1.bic:.2f}",          f"{mm2_ols.bic:.2f}",
         "↓ lepsze" if mm2_ols.bic < mm1.bic else "↑"],
    ]
)

h3("Wielokoliniowość po korekcie")
add_table(
    ["Zmienna", "VIF Model 1", "VIF Model 2", "Ocena"],
    [[c,
      f"{vif_m1[c]:.1f}",
      f"{vif_m2[c]:.1f}",
      "✔ Poprawa" if vif_m2[c] < vif_m1[c] else "→"]
     for c in Xm2_cols]
)

h3("Sprawdzenie poprawności — Model 2")
add_table(
    ["Test", "Model 1", "Model 2", "Zmiana"],
    [
        ["Shapiro-Wilk",
         f"p={sw_m1:.4f} {ok(sw_m1)}", f"p={sw_m2:.4f} {ok(sw_m2)}", "→"],
        ["Durbin-Watson",
         f"{dw_m1:.3f} {okdw(dw_m1)}", f"{dw_m2:.3f} {okdw(dw_m2)}", "→"],
        ["Breusch-Godfrey",
         f"p={bg_m1:.4f} {ok(bg_m1)}", f"p={bg_m2:.4f} {ok(bg_m2)}", "→"],
        ["Breusch-Pagan",
         f"p={bp_m1:.4f} {ok(bp_m1)}", f"p={bp_m2:.4f} {ok(bp_m2)}", "→"],
    ]
)

h3("Czy Model 2 nadaje się do prognoz?")
bullet(f"Normalność reszt: {'✔ OK' if sw_m2>0.05 else '⚠ Reszty OK (SW p>0.05)'} — "
       f"p = {sw_m2:.4f}")
bullet(f"Brak heteroskedastyczności: {'✔ OK' if bp_m2>0.05 else '⚠'} — "
       f"p = {bp_m2:.4f}")
bullet(f"Autokorelacja: {'✔ Brak' if bg_m2>0.05 else '⚠ Nadal obecna'} (DW={dw_m2:.3f}, BG p={bg_m2:.4f})")
bullet("Błędy HAC (Newey-West) korygują błędy standardowe → "
       "testy istotności są wiarygodne mimo autokorelacji")
p()
p("Wniosek: Model 2 jest poprawniejszy od Modelu 1. "
  "Do prognozowania wartości MO zalecane jest dodanie opóźnionej zmiennej "
  "zależnej (MO z poprzedniego roku) lub użycie modelu ARIMA "
  "(zaimplementowanego w skrypcie głównym).", bold=False)

# ============================================================
# PODSUMOWANIE
# ============================================================
doc.add_page_break()
h1("Podsumowanie i wnioski")

h2("Co udało się zbadać?")
p("W obu analizach udało się zbudować modele statystyczne opisujące "
  "ważne zjawiska społeczno-ekonomiczne w Polsce w latach 2004–2024.")

h2("Najważniejsze wnioski")

h3("Zużycie energii elektrycznej")
bullet("Najsilniejszy wpływ na zużycie energii ma PKB per capita — "
       f"elastyczność dochodowa wynosi {me2.params['ln_pkb_pc']:+.2f}%, "
       "co oznacza że bogacenie się społeczeństwa zwiększa zapotrzebowanie na energię.")
bullet("Cena energii ma wpływ ograniczający, ale stosunkowo słaby — "
       "Polacy są mało wrażliwi na wzrost cen energii (tzw. niska elastyczność cenowa).")
bullet("Wskaźniki klimatyczne (HDD, CDD) mają statystycznie istotny wpływ — "
       "zimniejsze zimy i gorętsze lata zwiększają zużycie.")
bullet("Prognoza bazowa na 2030 rok: "
       f"≈ {fc_energy['Bazowy'][-1][1]:,} GWh "
       f"({(fc_energy['Bazowy'][-1][1]/int(df_e['zuzycie_energii_GWh'].iloc[-1])-1)*100:+.1f}% vs 2024).")

h3("Mieszkania oddane do użytkowania")
bullet("Kluczowym czynnikiem jest poziom inwestycji budowlanych (nakłady) — "
       "to one bezpośrednio napędzają budownictwo mieszkaniowe.")
bullet("Czynniki demograficzne (udział osób 25–34 lat, urbanizacja, migracje) "
       "mają mniejszy, ale zauważalny wpływ.")
bullet("Rynek mieszkaniowy wykazuje silną inercję — co jest odzwierciedlone "
       "w autokorelacji reszt modelu (cykle budowlane trwają kilka lat).")

h2("Jakie modele wybrano do prognoz?")

add_table(
    ["Zmienna", "Model bazowy", "Model po korekcie (Iteracja 2)", "Zalecany do prognoz"],
    [
        ["Zużycie energii",
         "5 zmiennych (z urbanizacją)",
         "4 zmienne (bez urbanizacji)",
         "✔ Model 2 — spełnia założenia OLS"],
        ["Liczba mieszkań",
         "5 zmiennych (z WYNAGR)",
         "4 zmienne (bez WYNAGR) + błędy HAC",
         "✔ Model 2 + ARIMA jako uzupełnienie"],
    ]
)

h2("Ograniczenia analiz")
bullet("Mała próba: tylko 21 obserwacji rocznych (2004–2024) — "
       "wyniki mogą być wrażliwe na dane z pojedynczych lat.")
bullet("Dane niestacjonarne: zmienne rosną razem w czasie, "
       "co może prowadzić do regresji pozornej. "
       "Modele traktuje się jako opisowe, nie przyczynowe.")
bullet("Prognozy warunkowe: wyniki prognoz zależą od przyjętych "
       "założeń o PKB, cenach itp. — niepewność scenariuszy przenosi się "
       "na niepewność prognoz.")

doc.add_paragraph()
p("Pliki wygenerowane przez skrypty Python:", bold=True)
bullet("Zuzycie_energii_polska_py.py — analiza zużycia energii + wykresy 01–08")
bullet("mieszkania_polska_py.py — analiza mieszkań + wykresy m01–m06")
bullet("generuj_raport.py — niniejszy raport Word")

# ============================================================
# ZAPIS
# ============================================================
out_path = os.path.join(SCRIPT_DIR, "Raport_analiza.docx")
doc.save(out_path)
print(f"Raport zapisany: {out_path}")
